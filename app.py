"""
AI Internship Finder Agent — Flask Backend
==========================================
A Flask server that powers the AI Internship Finder chatbot.
Uses LangChain + LangGraph with OpenAI to create an intelligent agent that can
search for internships based on user skills, preferences, and resume.
"""

import os
import sys
# Force UTF-8 output so emoji in print() don't crash on Windows terminals
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
import json
import tempfile
import base64
import io
import requests
from flask import Flask, request, jsonify, render_template
from flask_cors import CORS
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from fpdf import FPDF
import docx as docx_lib

from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from langgraph.prebuilt import create_react_agent

from tools import search_internships
from mock_interview import mock_interview_bp

# ──────────────────────────────────────────────
# Load environment variables
# ──────────────────────────────────────────────
load_dotenv()

app = Flask(
    __name__,
    template_folder="internship_ai_agent",
    static_folder="internship_ai_agent",
    static_url_path="/static"
)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024  # 32 MB max upload

import uuid
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ── Simple file-based user store (email → user record) ──
USERS_FILE = os.path.join(os.path.dirname(__file__), 'uploads', 'users.json')

def _load_users():
    if os.path.exists(USERS_FILE):
        try:
            with open(USERS_FILE, 'r') as f:
                return json.load(f)
        except Exception:
            pass
    return {}

def _save_users(users):
    with open(USERS_FILE, 'w') as f:
        json.dump(users, f, indent=2)

def get_resume_stream_from_req(req):
    session_id = req.form.get("session_id") if req.form else None
    if not session_id and req.is_json:
        session_id = req.json.get("session_id")
    if session_id:
        filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
        if os.path.exists(filepath):
            return open(filepath, 'rb')
    if 'resume' in req.files:
        return req.files['resume'].stream
    return None

app.register_blueprint(mock_interview_bp)

# ──────────────────────────────────────────────
# Native Auth API Routes (no Node.js needed)
# ──────────────────────────────────────────────
import hashlib
import secrets

def _hash_password(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

def _make_token(user_id):
    return secrets.token_hex(32)


@app.route('/api/auth/register', methods=['POST'])
def auth_register():
    """Handle registration from the main register.html form (first_name + last_name fields)."""
    try:
        users = _load_users()
        # register.html sends form-data with first_name / last_name
        first_name = request.form.get('first_name', '').strip()
        last_name  = request.form.get('last_name', '').strip()
        email      = request.form.get('email', '').strip().lower()
        password   = request.form.get('password', '')
        target_role = request.form.get('role', '').strip()
        full_name  = (first_name + ' ' + last_name).strip() or email.split('@')[0]

        if not email or not password:
            return jsonify({'status': 'error', 'error': 'Email and password are required'}), 400
        if email in users:
            return jsonify({'status': 'error', 'error': 'User with this email already exists'}), 409

        # Handle optional resume upload
        skills = []
        role   = target_role or 'Software Engineer'
        session_id = None
        if 'resume' in request.files:
            resume_file = request.files['resume']
            session_id  = str(uuid.uuid4())
            filepath    = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
            resume_file.save(filepath)
            try:
                from mock_interview import extract_text_from_pdf, extract_skills, infer_role
                with open(filepath, 'rb') as f:
                    text = extract_text_from_pdf(f)
                skills = extract_skills(text)
                if skills:
                    role = infer_role(skills)
            except Exception as ex:
                print(f"[Register] Resume parse error: {ex}")

        user_id = str(uuid.uuid4())
        token   = _make_token(user_id)
        user_record = {
            '_id': user_id,
            'email': email,
            'fullName': full_name,
            'password_hash': _hash_password(password),
            'token': token,
            'extractedSkills': skills,
            'extractedRole': role,
            'session_id': session_id,
        }
        users[email] = user_record
        _save_users(users)

        # Store session_id under token key for dashboard compatibility
        return jsonify({'status': 'success', 'token': session_id or token,
                        'user_token': token, 'skills': skills, 'role': role}), 201
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'status': 'error', 'error': str(e)}), 500


@app.route('/api/auth/signup', methods=['POST'])
def auth_signup():
    try:
        users = _load_users()
        # Accept both JSON and form-data
        if request.is_json:
            data = request.get_json()
            email = data.get('email', '').strip().lower()
            password = data.get('password', '')
            full_name = data.get('fullName', data.get('full_name', ''))
        else:
            email = request.form.get('email', '').strip().lower()
            password = request.form.get('password', '')
            full_name = request.form.get('fullName', request.form.get('full_name', ''))

        if not email or not password:
            return jsonify({"error": "Email and password are required"}), 400
        if email in users:
            return jsonify({"error": "User with this email already exists"}), 409

        # Handle optional resume upload
        skills = []
        role = 'Software Engineer'
        session_id = None
        if 'resume' in request.files:
            resume_file = request.files['resume']
            session_id = str(uuid.uuid4())
            filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
            resume_file.save(filepath)
            # Extract skills from resume
            try:
                from mock_interview import extract_text_from_pdf, extract_skills, infer_role
                with open(filepath, 'rb') as f:
                    text = extract_text_from_pdf(f)
                skills = extract_skills(text)
                role = infer_role(skills) if skills else 'Software Engineer'
            except Exception as e:
                print(f"[Signup] Resume parsing error: {e}")

        user_id = str(uuid.uuid4())
        token = _make_token(user_id)
        user_record = {
            '_id': user_id,
            'email': email,
            'fullName': full_name or email.split('@')[0],
            'password_hash': _hash_password(password),
            'token': token,
            'extractedSkills': skills,
            'extractedRole': role,
            'session_id': session_id,
        }
        users[email] = user_record
        _save_users(users)

        safe_user = {k: v for k, v in user_record.items() if k != 'password_hash'}
        return jsonify({'token': token, 'user': safe_user, 'skills': skills, 'role': role}), 201
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/login', methods=['POST'])
def auth_login():
    try:
        users = _load_users()
        if request.is_json:
            data = request.get_json()
            email = data.get('email', '').strip().lower()
            password = data.get('password', '')
        else:
            email = request.form.get('email', '').strip().lower()
            password = request.form.get('password', '')

        if not email or not password:
            return jsonify({'error': 'Email and password are required'}), 400

        user = users.get(email)
        if not user:
            return jsonify({'error': 'Invalid email or password'}), 401

        # Support BOTH old plain-text 'password' field AND new 'password_hash' field
        old_plain = user.get('password')        # old format stored by old register
        new_hash  = user.get('password_hash')   # new hashed format

        password_ok = False
        if new_hash and new_hash == _hash_password(password):
            password_ok = True
        elif old_plain and old_plain == password:
            # Migrate: replace plain password with hash on first login
            password_ok = True
            user['password_hash'] = _hash_password(password)
            user.pop('password', None)

        if not password_ok:
            return jsonify({'error': 'Invalid email or password'}), 401

        # Ensure _id exists (old records may be missing it)
        if '_id' not in user:
            user['_id'] = str(uuid.uuid4())

        # Refresh token on every login
        token = _make_token(user['_id'])
        user['token'] = token
        users[email] = user
        _save_users(users)

        skills    = user.get('extractedSkills', [])
        role      = user.get('extractedRole', user.get('role', 'Software Engineer'))
        full_name = user.get('fullName') or (user.get('first_name', '') + ' ' + user.get('last_name', '')).strip() or email.split('@')[0]
        user['fullName'] = full_name

        safe_user = {k: v for k, v in user.items() if k not in ('password_hash', 'password')}
        return jsonify({'token': token, 'user': safe_user, 'skills': skills, 'role': role})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/upload-resume', methods=['POST'])
def auth_upload_resume():
    try:
        auth_header = request.headers.get('Authorization', '')
        token = auth_header.replace('Bearer ', '').strip()

        users = _load_users()
        user = next((u for u in users.values() if u.get('token') == token), None)
        if not user:
            return jsonify({'error': 'Unauthorized'}), 401

        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400

        resume_file = request.files['resume']
        session_id = str(uuid.uuid4())
        filepath = os.path.join(UPLOAD_FOLDER, f"{session_id}.pdf")
        resume_file.save(filepath)

        skills = []
        role = 'Software Engineer'
        try:
            from mock_interview import extract_text_from_pdf, extract_skills, infer_role
            with open(filepath, 'rb') as f:
                text = extract_text_from_pdf(f)
            skills = extract_skills(text)
            role = infer_role(skills) if skills else 'Software Engineer'
        except Exception as e:
            print(f"[UploadResume] Error: {e}")

        user['extractedSkills'] = skills
        user['extractedRole'] = role
        user['session_id'] = session_id
        users[user['email']] = user
        _save_users(users)

        safe_user = {k: v for k, v in user.items() if k != 'password_hash'}
        return jsonify({'user': safe_user, 'skills': skills, 'role': role})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ── In-memory reset token store: {email: {code, expires_at}} ──
import time as _time
_reset_tokens = {}

@app.route('/forgot-password')
def forgot_password_page():
    return render_template('forgot-password.html')

@app.route('/api/auth/forgot-password', methods=['POST'])
def auth_forgot_password():
    try:
        email = request.form.get('email', '').strip().lower()
        if not email:
            return jsonify({'error': 'Email is required'}), 400

        users = _load_users()
        if email not in users:
            # Don't reveal whether email exists — but for UX in local mode we tell them
            return jsonify({'error': 'No account found with that email address.'}), 404

        # Generate a 6-digit OTP
        import random as _random
        code = str(_random.randint(100000, 999999))
        _reset_tokens[email] = {
            'code': code,
            'expires_at': _time.time() + 600  # 10 minutes
        }
        print(f"[Password Reset] Code for {email}: {code}")

        # Return the code directly (local/demo mode — no email server)
        return jsonify({'status': 'ok', 'code': code,
                        'message': 'Reset code generated. Copy it from the box above.'})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/auth/reset-password', methods=['POST'])
def auth_reset_password():
    try:
        email    = request.form.get('email', '').strip().lower()
        code     = request.form.get('code', '').strip()
        new_pw   = request.form.get('new_password', '')

        if not email or not code or not new_pw:
            return jsonify({'error': 'Email, code and new password are required.'}), 400
        if len(new_pw) < 6:
            return jsonify({'error': 'Password must be at least 6 characters.'}), 400

        token_data = _reset_tokens.get(email)
        if not token_data:
            return jsonify({'error': 'No reset code found. Please request a new one.'}), 400
        if _time.time() > token_data['expires_at']:
            _reset_tokens.pop(email, None)
            return jsonify({'error': 'Reset code has expired. Please request a new one.'}), 400
        if token_data['code'] != code:
            return jsonify({'error': 'Invalid reset code. Check and try again.'}), 400

        # All good — update the password
        users = _load_users()
        if email not in users:
            return jsonify({'error': 'Account not found.'}), 404

        users[email]['password_hash'] = _hash_password(new_pw)
        users[email].pop('password', None)   # remove any old plain-text field
        _save_users(users)
        _reset_tokens.pop(email, None)       # invalidate used token

        return jsonify({'status': 'success', 'message': 'Password reset successfully.'})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ──────────────────────────────────────────────
# LangChain Agent Setup
# ──────────────────────────────────────────────

# ── LLM Setup: Google Gemini (free tier) primary, Ollama fallback ──
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
if GOOGLE_API_KEY:
    from langchain_google_genai import ChatGoogleGenerativeAI
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.0-flash",
        google_api_key=GOOGLE_API_KEY,
        temperature=0.3,
        convert_system_message_to_human=True
    )
    print("[OK] LLM: Google Gemini 2.0 Flash")
else:
    from langchain_ollama import ChatOllama
    llm = ChatOllama(model="llama3.2:1b", temperature=0.1)
    print("[WARN] GOOGLE_API_KEY not set - falling back to local Ollama llama3.2:1b")
    print("       Get a free Gemini key at: https://aistudio.google.com/")

# Tools available to the agent
tools = [search_internships]

# System prompt for the AI agent (Optimized for speed & anti-hallucination)
SYSTEM_PROMPT = """You are an AI Internship Finder Agent. Your goal is to help students find internships efficiently.

RULES:
1. ALWAYS use the `search_internships` tool to find internships.
2. The `search_internships` tool will return a fully formatted ```internship_cards``` block. YOU MUST ECHO THIS EXACT BLOCK TO THE USER VERBATIM!
3. DO NOT change the JSON. DO NOT change the apply_links. DO NOT invent URLs. DO NOT hallucinate.
4. Keep all conversational text extremely concise. Answer in 1 short sentence.
"""

# Create the agent using LangGraph's create_react_agent
agent = create_react_agent(llm, tools)

# ──────────────────────────────────────────────
# In-memory chat history (per session — resets on server restart)
# ──────────────────────────────────────────────
chat_histories = {}


def get_chat_history(session_id: str) -> list:
    """Get or create chat history for a session."""
    if session_id not in chat_histories:
        chat_histories[session_id] = [SystemMessage(content=SYSTEM_PROMPT)]
    return chat_histories[session_id]


# ──────────────────────────────────────────────
# Routes
# ──────────────────────────────────────────────

@app.route("/")
def index():
    """Serve the premium liquid glass landing page."""
    return render_template("landing.html")

@app.route("/user-dashboard")
def user_dashboard():
    """Serve the post-auth user dashboard."""
    return render_template("user-dashboard.html")

@app.route("/profile")
def profile_page():
    """Serve the user profile management page."""
    return render_template("profile.html")

@app.route("/dashboard")
def dashboard():
    """Serve the main chatbot hub."""
    return render_template("index.html")

@app.route("/bot")
def bot():
    """Serve the original interactive chatbot UI."""
    return render_template("bot.html")

@app.route('/login')
def login():
    return render_template('login.html')

@app.route('/register')
def register():
    return render_template('register.html')

# ── Auth routes are now proxied to Node.js backend via /api/auth/<path:subpath> ──

@app.route('/api/dashboard/analyze-resume', methods=['POST'])
def analyze_resume_for_dashboard():
    """
    Full pipeline:
      1. Extract text from PDF using pdfminer.six / PyPDF2 / pytesseract OCR
      2. Extract skills from text using the curated SKILLS_DB
      3. Scrape real internship listings from the internet for each top skill
      4. Score each listing against user's skill profile
      5. Return stats + top matched internships
    """
    from mock_interview import extract_text_from_pdf, extract_skills, infer_role
    import requests
    from bs4 import BeautifulSoup

    stream = get_resume_stream_from_req(request)
    if not stream:
        return jsonify({"error": "No resume found. Please provide session_id or file."}), 400

    try:
        # ── Step 1: Extract text ──────────────────────────────────────
        try:
            text = extract_text_from_pdf(stream)
        except Exception as e:
            return jsonify({"error": f"Text extraction failed: {str(e)}"}), 500

        if not text or len(text.strip()) < 30:
            return jsonify({"error": "Could not extract readable text from this PDF. Please try a text-based or clearer scan."}), 400

        # ── Step 2: Extract skills + infer role ──────────────────────
        skills = extract_skills(text)
        if not skills:
            skills = ['Python', 'Communication', 'Problem Solving']

        inferred_role = infer_role(skills)

        # ── Step 3: Scrape real internships ──────────────────────────
        def scrape_internships(skill_query, role_query, max_results=20):
            """Scrape real internship listings using DuckDuckGo HTML search."""
            jobs = []
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0 Safari/537.36'
            }

            search_terms = [
                f"{skill_query} internship 2025 2026",
                f"{role_query} intern site:linkedin.com OR site:internshala.com OR site:indeed.com OR site:glassdoor.com",
            ]

            for query in search_terms[:1]:
                try:
                    search_url = f"https://html.duckduckgo.com/html/?q={requests.utils.quote(query)}"
                    resp = requests.get(search_url, headers=headers, timeout=8)
                    soup = BeautifulSoup(resp.text, 'html.parser')

                    for result in soup.select('.result')[:max_results]:
                        title_el = result.select_one('.result__title')
                        snippet_el = result.select_one('.result__snippet')
                        url_el = result.select_one('.result__url')

                        if title_el and snippet_el:
                            title = title_el.get_text(strip=True)
                            snippet = snippet_el.get_text(strip=True)
                            url = url_el.get_text(strip=True) if url_el else ''

                            # Only include actual internship listings
                            combined = (title + ' ' + snippet).lower()
                            if any(w in combined for w in ['intern', 'internship', 'entry level', 'fresher', 'graduate']):
                                jobs.append({'title': title, 'snippet': snippet, 'url': url})

                except Exception as e:
                    print(f"[Scraper] Error: {e}")

            return jobs

        # ── Step 4: Score each scraped job against skills ─────────────
        def compute_match_score(job_text, user_skills_lower):
            """Returns a match percentage (0-100) based on skill overlap."""
            job_lower = job_text.lower()
            matched = sum(1 for skill in user_skills_lower if skill.lower() in job_lower)
            base_score = min(100, 60 + int((matched / max(len(user_skills_lower), 1)) * 40))
            return base_score

        top_skills_str = " ".join(skills[:5])
        raw_jobs = scrape_internships(top_skills_str, inferred_role, max_results=20)

        user_skills_lower = [s.lower() for s in skills]

        # Score and sort
        scored_jobs = []
        seen_titles = set()
        for job in raw_jobs:
            combined_text = job['title'] + ' ' + job['snippet']
            score = compute_match_score(combined_text, user_skills_lower)

            # Give bonus points for matching role keywords in title
            if inferred_role.lower().split()[0] in job['title'].lower():
                score = min(100, score + 8)

            clean_title = job['title'][:60]
            if clean_title not in seen_titles:
                seen_titles.add(clean_title)
                scored_jobs.append({
                    'title': clean_title,
                    'snippet': job['snippet'][:120],
                    'url': 'https://' + job['url'] if job['url'] and not job['url'].startswith('http') else job['url'],
                    'matchScore': score,
                    'company': _extract_company(job['url'], job['snippet']),
                    'mode': _guess_mode(job['snippet']),
                })

        scored_jobs.sort(key=lambda j: j['matchScore'], reverse=True)
        top_jobs = scored_jobs[:10]

        # ── Step 5: Compute dashboard stats ───────────────────────────
        matched_count = len(scored_jobs)
        avg_score = round(sum(j['matchScore'] for j in top_jobs) / max(len(top_jobs), 1)) if top_jobs else 75
        # Interviews completed is estimated based on skill breadth (simulated)
        interviews_completed = max(1, min(20, len(skills) // 3))

        return jsonify({
            "skills": skills,
            "role": inferred_role,
            "stats": {
                "matchedInternships": matched_count,
                "interviewsCompleted": interviews_completed,
                "avgMatchScore": avg_score,
                "dayStreak": max(1, len(skills) // 4)
            },
            "topMatches": top_jobs,
            "rawTextPreview": text[:500] + ("..." if len(text) > 500 else "")
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Backend crash during analysis: {str(e)}"}), 500


def _extract_company(url, snippet):
    """Try to extract company name from URL or snippet."""
    if 'linkedin.com' in url:
        return 'LinkedIn Listing'
    if 'indeed.com' in url:
        return 'Indeed Listing'
    if 'internshala.com' in url:
        return 'Internshala'
    if 'glassdoor.com' in url:
        return 'Glassdoor'
    # Try to pull from snippet
    for word in ['at ', 'with ', '@ ']:
        if word in snippet.lower():
            idx = snippet.lower().index(word) + len(word)
            return snippet[idx:idx+20].split()[0].strip(',.') if idx < len(snippet) else 'Company'
    return 'Tech Company'


def _guess_mode(snippet):
    s = snippet.lower()
    if 'remote' in s:
        return 'Remote'
    if 'hybrid' in s:
        return 'Hybrid'
    return 'On-site'


@app.route("/api/generate-cover-letter", methods=["POST"])
def generate_cover_letter():
    try:
        job_title = request.form.get("jobTitle", "Software Engineer")
        company = request.form.get("company", "the Company")
        from datetime import date
        today = date.today().strftime("%B %d, %Y")

        stream = get_resume_stream_from_req(request)
        if not stream:
            return jsonify({"error": "No resume found."}), 400

        from mock_interview import extract_text_from_pdf
        try:
            resume_text = extract_text_from_pdf(stream)
        except Exception as e:
            return jsonify({"error": f"Failed to read resume: {str(e)}"}), 500

        if len(resume_text.strip()) < 30:
            return jsonify({"error": "Could not extract readable text from this file."}), 400

        # Build strict official-format prompt
        prompt_text = f"""You are a professional cover letter writer. Generate a complete, formal business cover letter.

Requirements:
- Use the exact date: {today}
- Target position: {job_title}
- Target company: {company}
- Use the candidate's details extracted from the resume below.

Resume Content:
---
{resume_text[:2500]}
---

The cover letter MUST follow this EXACT official format:

[Candidate Full Name]
[Candidate Email]
[Candidate Phone, if found in resume]
[City, if found in resume]
{today}

Hiring Manager
{company}
Subject: Application for the Position of {job_title}

Dear Hiring Manager,

[Opening paragraph: Express strong enthusiasm for the role and company. Mention where you learned about the position.]

[Body paragraph: Highlight 2-3 specific skills, experiences, or projects from the resume most relevant to {job_title}. Use concrete details.]

[Closing paragraph: Reiterate interest, mention availability for an interview, and thank them for their consideration.]

Sincerely,
[Candidate Full Name]

Return ONLY the formatted cover letter text, no extra explanations."""

        # Call Ollama LLM
        response = llm.invoke([HumanMessage(content=prompt_text)])
        cl_text = response.content.strip()

        # Return plain text for editing
        return jsonify({"text": cl_text})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Failed to generate cover letter: {str(e)}"}), 500


@app.route("/api/download-cover-letter", methods=["POST"])
def download_cover_letter():
    """Generate PDF or DOCX from user-edited cover letter text."""
    try:
        data = request.json
        text = data.get("text", "").strip()
        fmt = data.get("format", "pdf").lower()

        if not text:
            return jsonify({"error": "No text provided"}), 400

        if fmt == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("helvetica", size=11)

            # Encode the whole text safely to latin-1, then write in one shot
            safe_text = text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 6, safe_text)

            pdf_bytes = bytes(pdf.output())
            pdf_b64 = base64.b64encode(pdf_bytes).decode('utf-8')
            return jsonify({"file_base64": pdf_b64, "format": "pdf"})

        elif fmt == "docx":
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = docx_lib.Document()

            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.2)
                section.right_margin = Inches(1.2)

            lines = text.split('\n')
            for i, line in enumerate(lines):
                if not line.strip():
                    doc.add_paragraph()
                elif i < 6:  # Header block — bold
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip())
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    p = doc.add_paragraph(line.strip())
                    p.runs[0].font.size = Pt(11) if p.runs else None

            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)
            docx_b64 = base64.b64encode(docx_io.read()).decode('utf-8')
            return jsonify({"file_base64": docx_b64, "format": "docx"})

        else:
            return jsonify({"error": "Invalid format. Use 'pdf' or 'docx'"}), 400

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Download generation failed: {str(e)}"}), 500


@app.route("/profile")
def profile():
    return render_template("profile.html")


@app.route("/contact")
def contact():
    return render_template("contact.html")

@app.route("/mock-interview/", defaults={"path": ""})
@app.route("/mock-interview/<path:path>")
def serve_mock_interview(path):
    return render_template("mock-interview.html")

@app.route("/documentation")
def documentation():
    return render_template("documentation.html")


@app.route("/privacy")
def privacy():
    return render_template("privacy.html")


@app.route("/scan", methods=["POST"])
def scan_jobs():
    """API endpoint for the 'Start AI Scan' button — returns real job results."""
    import random
    from tools import (scrape_linkedin, scrape_indeed, scrape_remotive,
                       scrape_freshershub, scrape_internshiphub, scrape_placementindia,
                       scrape_unstop, scrape_social_media)

    data = request.get_json() or {}
    query = data.get("query", "software intern India")

    all_cards = []
    seen_links = set()
    scrapers = [
        scrape_linkedin, scrape_indeed, scrape_remotive,
        scrape_freshershub, scrape_internshiphub, scrape_placementindia,
        scrape_unstop, scrape_social_media
    ]
    for func in scrapers:
        try:
            for card in func(query, days_ago=7):
                link = card.get("apply_link", "")
                if link and link not in seen_links:
                    seen_links.add(link)
                    all_cards.append(card)
        except Exception as e:
            print(f"[scan] scraper error: {e}")

    # Sort by applicants ascending (lowest competition first)
    all_cards.sort(key=lambda x: x.get("applicants", 10000))
    for i, card in enumerate(all_cards):
        card["id"] = i + 1

    return jsonify({"jobs": all_cards})


@app.route("/chat", methods=["POST"])
def chat():
    """Handle chat messages from the user."""
    try:
        data = request.get_json()
        user_message = data.get("message", "").strip()
        session_id = data.get("session_id", "default")

        if not user_message:
            return jsonify({"error": "Empty message"}), 400

        # Get chat history
        history = get_chat_history(session_id)

        # --- LIGHTNING-FAST TOOL INTERCEPTOR ---
        # Don't wait 15-30 seconds for the local CPU LLM to process if we already know they want jobs!
        search_keywords = ["intern", "job", "role", "developer", "engineer", "position", "work", "find", "show", "remote"]
        needs_search = any(k in user_message.lower() for k in search_keywords)
        
        ai_response = ""
        
        if needs_search:
            # ── SOURCE FILTER: detect if the user wants a specific platform only ──
            # Determine source filter
            msg_lower = user_message.lower()
            discovered_sources = []
            
            source_map = {
                "linkedin": "linkedin",
                "indeed": "indeed",
                "remotive": "remotive",
                "freshershub": "freshershub",
                "placementhub": "placementhub",
                "internshiphub": "internshiphub",
                "unstop": "unstop",
                "twitter": "social",
                "facebook": "social",
                "instagram": "social",
                "social": "social",
                "placementindia": "placementindia",
            }
            
            for kw, src in source_map.items():
                if kw in msg_lower:
                    discovered_sources.append(src)
            # Deduplicate (e.g. twitter + facebook both map to "social")
            discovered_sources = list(dict.fromkeys(discovered_sources))
                    
            sources_param = ",".join(discovered_sources) if discovered_sources else "linkedin,indeed,remotive,freshershub,placementhub,internshiphub,unstop,social"

            tool_output = search_internships.invoke({
                "query": user_message,
                "sources": sources_param
            })
            
            source_labels = discovered_sources if discovered_sources else ["All Sources"]
            source_title = "/".join(s.capitalize() for s in source_labels)
            
            if "```internship_cards" in tool_output:
                ai_response = f"Here are live jobs from **{source_title}**:\n\n" + tool_output
            else:
                ai_response = tool_output
        else:
            # Only run the LLM natively if it's conversational small talk
            llm_with_tools = llm.bind_tools(tools)
            ai_msg = llm_with_tools.invoke(history)
            
            if ai_msg.tool_calls:
                query_arg = ai_msg.tool_calls[0]["args"].get("query", user_message)
                tool_output = search_internships.invoke({"query": query_arg, "sources": "linkedin,indeed,remotive"})
                if "```internship_cards" in tool_output:
                    ai_response = "Here are the live job listings:\n\n" + tool_output
                else:
                    ai_response = tool_output
            else:
                ai_response = ai_msg.content
                if "{" in ai_response and "}" in ai_response:
                    ai_response = "I couldn't quite understand that. Are you looking for a specific internship?"

        if not ai_response:
            ai_response = "I'm sorry, I couldn't process that. Could you try rephrasing? 😊"

        # Update history with AI's tool or text response
        if "```internship_cards" in ai_response:
            # Do NOT store 3,000 tokens of JSON in the LLM's chat history! It will freeze local models.
            history.append(AIMessage(content="I successfully found and displayed the matching internship cards to the user."))
        else:
            history.append(AIMessage(content=ai_response))

        # Keep history extremely lightweight to ensure local LLM runs fast (System msg + last 4 messages)
        if len(history) > 5:
            history[:] = [history[0]] + history[-4:]

        return jsonify({
            "response": ai_response,
            "status": "success"
        })

    except Exception as e:
        print(f"Error in /chat: {e}")
        error_msg = str(e)
        if "connection error" in error_msg.lower() or "connectionrefusederror" in error_msg.lower():
            return jsonify({
                "response": "⚠️ **Ollama Error**: Could not connect to local Ollama instance. Please ensure Ollama is running and you have run `ollama run mistral`.",
                "status": "success"
            })
        return jsonify({
            "error": f"Something went wrong: {str(e)}",
            "status": "error"
        }), 500


@app.route("/upload-resume", methods=["POST"])
def upload_resume():
    """
    Handle PDF resume uploads:
    1. Extract text via PyPDF2
    2. Parse key skills with a fast local regex extractor (NO LLM = instant)
    3. Search all platforms for jobs posted in the LAST 7 DAYS
    4. Return job cards + a skill summary
    """
    import re

    try:
        stream = get_resume_stream_from_req(request)
        if not stream:
            return jsonify({"error": "No resume found"}), 400

        session_id = request.form.get("session_id", "default")

        # ── Step 1: Extract text from PDF ─────────────────────────────────
        reader = PdfReader(stream)
        resume_text = ""
        for page in reader.pages:
            text = page.extract_text()
            if text:
                resume_text += text + "\n"

        if not resume_text.strip():
            return jsonify({"error": "Could not extract text from the PDF. Please try a text-based PDF."}), 400

        # ── Step 2: Fast keyword extractor (NO LLM — instant) ─────────────
        text_lower = resume_text.lower()

        TECH_SKILLS = [
            "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
            "react", "angular", "vue", "node.js", "nodejs", "django", "flask", "fastapi",
            "spring", "express", "next.js", "nextjs", "sql", "mysql", "postgresql",
            "mongodb", "redis", "machine learning", "deep learning", "nlp",
            "computer vision", "data science", "tensorflow", "pytorch", "keras",
            "scikit-learn", "pandas", "numpy", "aws", "azure", "gcp", "docker",
            "kubernetes", "git", "html", "css", "rest api", "graphql",
            "android", "ios", "flutter", "react native", "swift", "kotlin",
            "blockchain", "solidity", "web3", "devops"
        ]

        ROLE_KEYWORDS = [
            "software engineer", "software developer", "backend developer",
            "frontend developer", "full stack developer", "data scientist",
            "data analyst", "ml engineer", "machine learning engineer",
            "ai engineer", "devops engineer", "android developer", "ios developer",
            "mobile developer", "web developer", "intern", "fresher"
        ]

        found_skills = [s for s in TECH_SKILLS if s in text_lower]
        found_roles = [r for r in ROLE_KEYWORDS if r in text_lower]

        # ── Step 3: Build ROTATING query variations ────────────────────────
        # Shuffle skills so every upload picks different combos as primary keys
        import random
        shuffled_skills = found_skills[:] if found_skills else ["software"]
        random.shuffle(shuffled_skills)

        primary_role = found_roles[0] if found_roles else "developer"

        # Build up to 3 unique queries from different skill combos
        queries = []
        if len(shuffled_skills) >= 2:
            queries.append(f"{shuffled_skills[0]} {primary_role} intern")
            queries.append(f"{shuffled_skills[1]} developer intern india")
        if len(shuffled_skills) >= 3:
            queries.append(f"{shuffled_skills[2]} engineer internship")
        if not queries:
            queries = [f"{shuffled_skills[0]} {primary_role} intern"]

        search_query = queries[0]  # primary query for display

        # ── Step 4: Multi-query scraping with STRICT 7-day filter ──────────
        from tools import scrape_linkedin, scrape_indeed, scrape_remotive, \
                          scrape_freshershub, scrape_internshiphub, scrape_placementindia, \
                          scrape_unstop, scrape_social_media

        seen_links = set()
        all_cards = []

        # Run all generated queries to get broad variety
        scrapers = [
            scrape_linkedin, scrape_indeed, scrape_remotive, 
            scrape_freshershub, scrape_internshiphub, scrape_placementindia,
            scrape_unstop, scrape_social_media
        ]

        for q in queries:
            for scraper_func in scrapers:
                try:
                    # Fetch fresh jobs (7 days) for this query variation
                    results = scraper_func(q, days_ago=7)
                    for card in results:
                        link = card.get("apply_link", "")
                        if link and link not in seen_links:
                            seen_links.add(link)
                            all_cards.append(card)
                except Exception as e:
                    print(f"Scraper error for query '{q}': {e}")

        # Shuffle the final pool to ensure a different visual order every time
        random.shuffle(all_cards)

        # Re-number and badge (after shuffle)
        for i, card in enumerate(all_cards):
            card["id"] = i + 1
            source = card.pop("source", "Job Board")
            card["title"] = f"[{source}] {card['title']}"

        # ── Step 4: Build the response ─────────────────────────────────────
        skills_display = ", ".join(f"**{s.title()}**" for s in found_skills[:8]) or "General Software Skills"
        roles_display = ", ".join(r.title() for r in found_roles[:3]) or "Software Developer"

        if all_cards:
            cards_block = "```internship_cards\n" + json.dumps(all_cards, indent=2) + "\n```"
            intro = (
                f"📄 **Resume Analyzed!**\n\n"
                f"🔑 **Skills Detected:** {skills_display}\n"
                f"💼 **Target Roles:** {roles_display}\n"
                f"🔍 **Searched for:** `{search_query}`\n"
                f"⏰ **Filter:** Jobs posted in the **last 7 days** only\n\n"
                f"Found **{len(all_cards)} fresh, open opportunities** matching your profile:\n\n"
            )
            ai_response = intro + cards_block
        else:
            ai_response = (
                f"📄 **Resume Analyzed!**\n\n"
                f"🔑 **Skills Detected:** {skills_display}\n"
                f"💼 **Target Roles:** {roles_display}\n\n"
                f"⚠️ No jobs found in the **last 7 days** for `{search_query}`. "
                f"Ask me to search with a wider range, e.g. *'find python developer jobs from last month'*."
            )

        history = get_chat_history(session_id)
        history.append(HumanMessage(content=f"[Resume analyzed — skills: {', '.join(found_skills[:5])}]"))
        history.append(AIMessage(content=ai_response))
        if len(history) > 5:
            history[:] = [history[0]] + history[-4:]

        return jsonify({
            "response": ai_response,
            "status": "success",
            "skills_found": found_skills,
            "roles_found": found_roles,
            "search_query": search_query,
            "jobs_count": len(all_cards)
        })

    except Exception as e:
        print(f"Error in /upload-resume: {e}")
        return jsonify({"error": f"Failed to process resume: {str(e)}", "status": "error"}), 500



@app.route("/api/avatar/session", methods=["POST"])
def create_avatar_session():
    """Create a new Akool Live Avatar streaming session."""
    akool_api_key = os.getenv("AKOOL_API_KEY")
    if not akool_api_key or akool_api_key == "your_akool_api_key_here":
        return jsonify({"error": "AKOOL_API_KEY is not set in the .env file"}), 500

    url = "https://openapi.akool.com/api/v3/streamingAvatar/session/create"
    headers = {
        "Authorization": f"Bearer {akool_api_key}",
        "Content-Type": "application/json"
    }
    
    # User should set AKOOL_AVATAR_ID in .env, fallback to a commonly used default format or name
    avatar_id = os.getenv("AKOOL_AVATAR_ID", "default_avatar") 
    
    payload = {
        "avatar_id": avatar_id
    }

    try:
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()
        data = response.json()
        return jsonify(data)
    except requests.exceptions.HTTPError as e:
        print(f"HTTP Error creating Akool session: {e.response.text}")
        return jsonify({"error": "Failed to create avatar session", "details": e.response.text}), e.response.status_code
    except Exception as e:
        print(f"Error creating Akool session: {e}")
        return jsonify({"error": str(e)}), 500



# ──────────────────────────────────────────────
# Mock Interview Question Generator
# ──────────────────────────────────────────────
@app.route('/api/interview/generate-questions', methods=['POST'])
def generate_interview_questions():
    try:
        data = request.get_json() or {}
        skills = data.get('skills', [])
        difficulty = data.get('difficulty', 'Medium')
        count = int(data.get('count', 5))

        skills_str = ', '.join(skills) if skills else 'general software engineering'

        difficulty_guidance = {
            'Easy': 'simple, conceptual, beginner-friendly questions that test basic understanding and definitions',
            'Medium': 'practical, scenario-based questions that test applied knowledge and real-world usage',
            'Hard': 'complex, architectural, system-design and edge-case questions requiring deep expert-level reasoning'
        }.get(difficulty, 'practical questions')

        prompt = f"""You are a professional technical interviewer. Generate exactly {count} interview questions for a candidate with the following skills: {skills_str}.

Difficulty level: {difficulty} — Ask {difficulty_guidance}.

Rules:
- Exactly ONE of the {count} questions must be a "coding" question (e.g. "Write a function to...", "Implement a script that..."). The other {count-1} should be spoken/conceptual.
- Each question must be directly related to the candidate's listed skills.
- Questions must be distinct and not repeat concepts.
- Return ONLY a JSON array of {count} objects, with no outside text.
- Each object must have two keys: "text" (the question string) and "is_coding" (boolean).
- Example format: 
[
  {{"text": "Explain how event loops work.", "is_coding": false}},
  {{"text": "Write a Python function to reverse a string.", "is_coding": true}}
]

Output the JSON array now:"""

        response = llm.invoke([HumanMessage(content=prompt)])
        raw = response.content.strip()

        # Parse the JSON from the response
        import re
        match = re.search(r'\[\s*\{.*\}\s*\]', raw, re.DOTALL)
        
        questions = []
        if match:
            try:
                questions = json.loads(match.group())
            except:
                pass
                
        # Fallback if parsing failed
        if not questions or not isinstance(questions, list):
            questions = [
                {"text": f"Can you explain your experience with {skills[0] if skills else 'software development'}?", "is_coding": False},
                {"text": "Describe a challenging technical problem you solved recently.", "is_coding": False},
                {"text": f"Write a function to implement a basic feature using {skills[0] if skills else 'any language'}.", "is_coding": True},
                {"text": "How do you approach debugging a complex issue?", "is_coding": False},
                {"text": "What software development best practices do you follow?", "is_coding": False}
            ][:count]

        # Ensure we have exactly `count` items
        while len(questions) < count:
            questions.append({"text": "Where do you see yourself growing technically in the next year?", "is_coding": False})
        questions = questions[:count]

        return jsonify({'questions': questions, 'difficulty': difficulty, 'count': len(questions)})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/api/fast-extract-skills', methods=['POST'])
def fast_extract_skills():
    if 'resume' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    try:
        file = request.files['resume']
        from mock_interview import extract_text_from_pdf, extract_skills
        text = extract_text_from_pdf(file.stream)
        skills = extract_skills(text)
        return jsonify({"extractedSkills": skills})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ──────────────────────────────────────────────
# Run the server
# ──────────────────────────────────────────────
if __name__ == "__main__":
    print("[START] AI Internship Finder Agent is running!")
    print("[INFO] Open: http://127.0.0.1:5000")
    print("")
    app.run(debug=True, port=5000)
